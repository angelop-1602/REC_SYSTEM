import os
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
import win32com.client as win32

class DocumentGenerator:
    def __init__(self, input_file, output_dir):
        self.input_file = input_file
        self.output_dir = output_dir
        self.rec_code = ""
        self.protocol_title = ""
        self.principal_investigator = ""
        self.adviser = ""

    def generate_document(self, rec_code, protocol_title, principal_investigator, adviser, action="save"):
        """Generate the document and either save or save+print it based on action."""
        # Set inputs to object properties
        self.rec_code = rec_code
        self.protocol_title = protocol_title
        self.principal_investigator = principal_investigator
        self.adviser = adviser

        formatted_protocol_title = self.title_case(protocol_title)
        formatted_principal_investigator = self.title_case(principal_investigator)
        formatted_adviser = self.format_adviser_name_with_extension(adviser)

        # Calculate today's date and its formatted version
        today_date = datetime.now()
        formatted_today_date = today_date.strftime("%d %b. %Y")
        full_today_date = today_date.strftime("%d %B %Y")

        # Calculate LESSDATE
        less_date = today_date - timedelta(days=10)
        formatted_less_date = less_date.strftime("%d %b. %Y")

        date_plus = (today_date + timedelta(days=365)).strftime("%d %b. %Y")
        dateplus = (today_date + timedelta(days=21)).strftime("%d %b. %Y")

        # Create output directory based on today's date
        date_folder_name = today_date.strftime("%Y-%m-%d")
        output_date_dir = os.path.join(self.output_dir, date_folder_name)

        if not os.path.exists(output_date_dir):
            os.makedirs(output_date_dir)

        # Extract base name from input file
        base_input_file_name = os.path.splitext(os.path.basename(self.input_file))[0]  # Get input file name without extension
        output_file_name = f'{formatted_principal_investigator.replace(" ", "_")} ({base_input_file_name}).docx'  # Include input file name
        today_output_file_path = os.path.join(output_date_dir, output_file_name)

        # Check if a document has been created today
        if os.path.exists(today_output_file_path):
            base_name, ext = os.path.splitext(output_file_name)
            counter = 1
            while os.path.exists(today_output_file_path):
                today_output_file_path = os.path.join(output_date_dir, f'{base_name}_{counter}{ext}')
                counter += 1

        replacements = {
            "<<DATETODAY>>": full_today_date,
            "<<DATE_TODAY>>": formatted_today_date,
            "<<LESSDATE>>": formatted_less_date,
            "<<REC_CODE>>": rec_code,
            "<<PROTOCOL_TITLE>>": formatted_protocol_title,
            "<<PRINCIPAL_INVESTIGATOR>>": formatted_principal_investigator,
            "<<ADVISER>>": formatted_adviser,
            "<<DATE_PLUS>>": date_plus,
            "<<DATEPLUS>>": dateplus
        }

        try:
            self.replace_placeholders(self.input_file, replacements, today_output_file_path)
            print(f"Document saved as: {today_output_file_path}")

            # Decide based on action: Save or Save + Print
            if action == "save_and_print":
                self.print_word_document(today_output_file_path)

            # After saving or printing, clear the inputs
            self.reset_inputs()
        except Exception as e:
            print(f"An error occurred: {e}")

    def save_document(self, rec_code, protocol_title, principal_investigator, adviser):
        """Method for saving the document only."""
        self.generate_document(rec_code, protocol_title, principal_investigator, adviser, action="save")

    def save_and_print_document(self, rec_code, protocol_title, principal_investigator, adviser):
        """Method for saving and printing the document."""
        self.generate_document(rec_code, protocol_title, principal_investigator, adviser, action="save_and_print")

    def print_word_document(self, file_path):
        """Print the generated Word document."""
        try:
            word = win32.Dispatch("Word.Application")
            word.Visible = False  
            doc = word.Documents.Open(file_path)
            doc.PrintOut()  
            doc.Close(False)  
            word.Quit()  
            print("Document sent to printer.")
        except Exception as e:
            print(f"An error occurred while printing: {e}")

    def replace_placeholders(self, doc_path, replacements, output_path):
        doc = Document(doc_path)
        print(f"Loaded document from: {doc_path}")

        for para in doc.paragraphs:
            for key, value in replacements.items():
                if key in para.text:
                    print(f"Replacing '{key}' with '{value}' in paragraph: {para.text}")
                    self.replace_text_in_paragraph(para, key, value)

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in para.text:
                                print(f"Replacing '{key}' with '{value}' in cell: {para.text}")
                                self.replace_text_in_paragraph(para, key, value)

        doc.save(output_path)
        print(f"Document saved to: {output_path}")

    def replace_text_in_paragraph(self, para, key, value):
        runs = para.runs
        full_text = ''.join(run.text for run in runs)

        if key in full_text:
            for run in runs:
                para._element.remove(run._element)

            parts = full_text.split(key)
            for i, part in enumerate(parts):
                if i > 0:
                    run = para.add_run(value)
                    self.format_run(run, 'Times New Roman', 12, bold=True if key in ["<<REC_CODE>>", "<<DATETODAY>>"] else False, uppercase=True if key in ["<<REC_CODE>>", "<<DATETODAY>>"] else False)

                run = para.add_run(part)
                self.format_run(run, 'Times New Roman', 12)

    def format_run(self, run, font_name, font_size, bold=False, italic=False, underline=False, uppercase=False):
        font = run.font
        font.name = font_name
        font.size = Pt(font_size)
        font.bold = bold
        font.italic = italic
        font.underline = underline
        if uppercase:
            run.text = run.text.upper()

    def title_case(self, text):
        """Convert a string to title case, excluding specified words."""
        words = text.split()
        exceptions = {
            'and', 'or', 'the', 'in', 'a', 'an', 
            'of', 'for', 'to', 'but', 'nor', 'on', 
            'at', 'by', 'with', 'as', 'about', 'if', 
            'than', 'so', 'up', 'down', 'out', 'over', 
            'under', 'between', 'into', 'through'
        }
        
        title_cased_words = []
        for index, word in enumerate(words):
            if index == 0 or word.lower() not in exceptions:
                title_cased_words.append(word.capitalize())
            else:
                title_cased_words.append(word.lower())
        return ' '.join(title_cased_words)

    def format_adviser_name_with_extension(self, adviser_name):
        """Convert any recognized extensions in the adviser's name to uppercase and capitalize each word."""
        adviser_name = adviser_name.title()
        parts = adviser_name.split(',')
        base_name = parts[0].strip()

        extensions = [part.strip().upper() for part in parts[1:]]
        if extensions:
            return f"{base_name}, " + ', '.join(extensions)

        return adviser_name

    def reset_inputs(self):
        """Reset the input fields to empty after document creation and printing."""
        self.rec_code = ""
        self.protocol_title = ""
        self.principal_investigator = ""
        self.adviser = ""
        print("Input fields have been cleared.")
